from abc import ABC, abstractmethod
import html
import re
import tempfile

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    PageBreak,
    Spacer,
)
from reportlab.lib.enums import TA_JUSTIFY

# ---------------------------------------------------------------------------
# Try to register a CJK-capable font so Japanese / Chinese output works.
# If the Noto font isn't installed we fall back to Helvetica (Latin-only).
# ---------------------------------------------------------------------------
_BODY_FONT = "Helvetica"

try:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    # Common locations where Noto Sans CJK ships on Linux / Docker images
    _NOTO_CANDIDATES = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/noto/NotoSans-Regular.ttf",
    ]

    for path in _NOTO_CANDIDATES:
        if os.path.isfile(path):
            pdfmetrics.registerFont(TTFont("NotoSans", path))
            _BODY_FONT = "NotoSans"
            break
except Exception:
    pass  # Stick with Helvetica


class Writer(ABC):
    @abstractmethod
    def write(self, pages: list[str], original_path: str) -> str:
        pass


class PdfWriter(Writer):
    """Produce a nicely formatted PDF with proper text wrapping."""

    # Page geometry
    _MARGIN = 25 * mm  # ~1 inch

    def write(self, pages: list[str], original_path: str) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            file_path = tmp.name

        doc = SimpleDocTemplate(
            file_path,
            pagesize=A4,
            leftMargin=self._MARGIN,
            rightMargin=self._MARGIN,
            topMargin=self._MARGIN,
            bottomMargin=self._MARGIN,
            allowSplitting=1,    # Allow paragraph splitting across pages
            autoNextPageTemplate=0,
        )

        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "TranslatedBody",
            parent=styles["Normal"],
            fontName=_BODY_FONT,
            fontSize=11,
            leading=15,          # line spacing
            alignment=TA_JUSTIFY,  # Force full justification
            spaceAfter=12,       # space between paragraphs
            spaceBefore=0,       # no space before paragraphs
            leftIndent=0,        # no left indent
            rightIndent=0,       # no right indent
            firstLineIndent=0,   # no first line indent
            wordWrap='LTR',      # left-to-right word wrapping
            splitLongWords=1,    # Allow word breaking
            allowWidows=0,       # Prevent widow lines
            allowOrphans=0,      # Prevent orphan lines
            adjustFontSize=1,    # Allow slight font adjustment for better fit
            hyphenationLang='en_US',  # Enable hyphenation
        )

        story: list = []

        for page_idx, page_text in enumerate(pages):
            if page_idx > 0:
                story.append(PageBreak())

            # Split on double-newlines → paragraphs
            paragraphs = page_text.split("\n\n")
            for para in paragraphs:
                para = para.strip()
                if not para:
                    continue

                # Clean up the paragraph for better justification
                # Remove single newlines and extra whitespace
                para = para.replace("\n", " ")
                # Normalize multiple spaces to single spaces
                para = re.sub(r'\s+', ' ', para)
                para = para.strip()

                # Skip very short paragraphs that might not justify well
                if len(para) < 10:
                    continue

                # Add soft hyphens for better word breaking
                # This helps with justification of long words
                para = re.sub(r'([a-zA-Z]{6,})', '\\1\u00AD', para)

                # Escape HTML entities so angle-brackets in content survive
                safe = html.escape(para)

                # Create paragraph with explicit style to ensure justification
                p = Paragraph(safe, body_style)
                story.append(p)

        if not story:
            story.append(Paragraph("&nbsp;", body_style))

        doc.build(story)
        return file_path

class DocxPlainWriter(Writer):
    """Create a simple DOCX when structured formatting is unavailable."""

    def write(self, pages: list[str], original_path: str) -> str:
        doc = Document()

        for page_idx, page_text in enumerate(pages):
            paragraphs = page_text.split("\n\n") if page_text else []

            if not paragraphs:
                doc.add_paragraph("")
            else:
                for para in paragraphs:
                    para = para.strip()
                    if not para:
                        continue
                    doc.add_paragraph(para)

            if page_idx < len(pages) - 1:
                doc.add_page_break()

        if not doc.paragraphs:
            doc.add_paragraph("(empty document)")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            output_path = tmp.name

        doc.save(output_path)
        return output_path


class WriterFactory:
    @staticmethod
    def get_writer(file_path: str, output_format: str = "PDF") -> Writer:
        file_path_lower = file_path.lower()
        fmt = (output_format or "PDF").upper()

        if fmt == "DOCX":
            return DocxPlainWriter()

        # For PPTX files, we need special handling to preserve formatting
        # The orchestrator should handle PPTX separately, but this provides a fallback
        if file_path_lower.endswith(".pptx"):
            # Return PdfWriter as fallback (convert to PDF)
            # The orchestrator should use PptxWriter directly for format preservation
            return PdfWriter()
        else:
            # All other documents are translated to PDF format
            # regardless of input format (pdf, docx, odt, txt)
            return PdfWriter()
