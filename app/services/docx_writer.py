"""
DOCX Writer - Preserves document structure, formatting, and metadata
"""
import tempfile
from copy import deepcopy
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


class DocxWriter:
    """
    Writes translated text back to DOCX while preserving:
    - Original document structure
    - Paragraph formatting and styles
    - Text runs with formatting (bold, italic, font, size, color)
    - Tables with structure
    - Images
    """

    def write(self, document_data: Dict[str, Any], original_path: str) -> str:
        """
        Write translated document back to DOCX format.

        Args:
            document_data: Dict with translated elements
            original_path: Path to original DOCX (used as template)

        Returns:
            Path to the output DOCX file
        """
        doc = Document(original_path)
        body = doc.element.body

        # Preserve original ordering (including media) by cloning elements
        original_elements = [deepcopy(element) for element in body]

        # Clear body content so we can rebuild deterministically
        for element in list(body):
            body.remove(element)

        elements = document_data.get('elements', [])
        elem_idx = 0

        for original_element in original_elements:
            if isinstance(original_element, CT_P):
                if self._paragraph_has_text(original_element):
                    if elem_idx < len(elements) and elements[elem_idx].get('type') == 'paragraph':
                        media_runs = self._extract_media_runs(original_element)
                        self._add_paragraph(doc, elements[elem_idx], media_runs)
                        elem_idx += 1
                    else:
                        body.append(original_element)
                else:
                    body.append(original_element)
            elif isinstance(original_element, CT_Tbl):
                if elem_idx < len(elements) and elements[elem_idx].get('type') == 'table':
                    self._add_table(doc, elements[elem_idx])
                    elem_idx += 1
                else:
                    body.append(original_element)
            else:
                # Non-text elements (pictures, drawings, etc.) are appended as-is
                body.append(original_element)

        # Append any leftover translated elements (should rarely happen)
        while elem_idx < len(elements):
            element = elements[elem_idx]
            if element.get('type') == 'paragraph':
                self._add_paragraph(doc, element)
            elif element.get('type') == 'table':
                self._add_table(doc, element)
            elem_idx += 1

        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            output_path = tmp.name

        doc.save(output_path)
        return output_path

    def _add_paragraph(self, doc: Document, para_data: Dict[str, Any], media_runs: List | None = None):
        """Add a paragraph with formatting."""
        paragraph = doc.add_paragraph()

        # Set paragraph style if available
        style_name = para_data.get('style')
        if style_name:
            try:
                paragraph.style = style_name
            except:
                pass  # Style might not exist, skip

        # Set alignment
        alignment = para_data.get('alignment')
        if alignment is not None:
            paragraph.alignment = alignment

        # Add runs with formatting
        runs_data = para_data.get('runs', [])
        if runs_data:
            for run_data in runs_data:
                run = paragraph.add_run(run_data.get('text', ''))

                # Apply formatting
                if run_data.get('bold') is not None:
                    run.bold = run_data['bold']
                if run_data.get('italic') is not None:
                    run.italic = run_data['italic']
                if run_data.get('underline') is not None:
                    run.underline = run_data['underline']

                # Font name
                if run_data.get('font_name'):
                    run.font.name = run_data['font_name']

                # Font size
                if run_data.get('font_size'):
                    run.font.size = Pt(run_data['font_size'])

                # Font color
                if run_data.get('font_color'):
                    try:
                        color_str = run_data['font_color']
                        # Parse RGB color (format: "RRGGBB")
                        if len(color_str) >= 6:
                            r = int(color_str[0:2], 16)
                            g = int(color_str[2:4], 16)
                            b = int(color_str[4:6], 16)
                            run.font.color.rgb = RGBColor(r, g, b)
                    except:
                        pass  # Skip if color parsing fails
        else:
            paragraph.add_run(para_data.get('text', ''))

        if media_runs:
            for media_run in media_runs:
                paragraph._p.append(media_run)

    def _add_table(self, doc: Document, table_data: Dict[str, Any]):
        """Add a table with structure and formatting."""
        num_rows = table_data.get('num_rows', 0)
        num_cols = table_data.get('num_cols', 0)

        if num_rows == 0 or num_cols == 0:
            return

        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'  # Basic grid style

        # Fill table cells
        rows_data = table_data.get('rows', [])

        for row_idx, row_data in enumerate(rows_data):
            if row_idx >= num_rows:
                break

            for col_idx, cell_data in enumerate(row_data):
                if col_idx >= num_cols:
                    break

                cell = table.rows[row_idx].cells[col_idx]

                # Clear default paragraph
                cell.paragraphs[0].clear()

                # Add cell paragraphs with formatting
                cell_paragraphs = cell_data.get('paragraphs', [])

                if cell_paragraphs:
                    for para_idx, para_data in enumerate(cell_paragraphs):
                        if para_idx == 0:
                            # Use existing paragraph
                            paragraph = cell.paragraphs[0]
                        else:
                            # Add new paragraph
                            paragraph = cell.add_paragraph()

                        # Add runs with formatting
                        runs_data = para_data.get('runs', [])
                        if runs_data:
                            for run_data in runs_data:
                                run = paragraph.add_run(run_data.get('text', ''))

                                # Apply formatting
                                if run_data.get('bold') is not None:
                                    run.bold = run_data['bold']
                                if run_data.get('italic') is not None:
                                    run.italic = run_data['italic']
                                if run_data.get('underline') is not None:
                                    run.underline = run_data['underline']
                                if run_data.get('font_name'):
                                    run.font.name = run_data['font_name']
                                if run_data.get('font_size'):
                                    run.font.size = Pt(run_data['font_size'])
                        else:
                            paragraph.add_run(para_data.get('text', ''))
                else:
                    # Fallback: just add cell text
                    cell.text = cell_data.get('text', '')

    def _paragraph_has_text(self, paragraph_element: CT_P) -> bool:
        """Check if a paragraph originally contained text (matches loader logic)."""
        text_nodes = paragraph_element.xpath('.//w:t')
        return any((t.text or '').strip() for t in text_nodes)

    def _extract_media_runs(self, paragraph_element: CT_P) -> List:
        """Clone runs containing drawings or pictures so they can be re-attached."""
        media_runs: List = []

        for run in paragraph_element.xpath('.//w:r'):
            has_drawing = bool(run.xpath('.//w:drawing'))
            has_picture = bool(run.xpath('.//w:pict'))

            if has_drawing or has_picture:
                media_runs.append(deepcopy(run))

        return media_runs
